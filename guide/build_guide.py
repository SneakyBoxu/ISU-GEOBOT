"""
Build the simple how-it-works guide.

    python guide/build_guide.py

Organised smallest-first: what each piece IS, then how each mechanism WORKS,
then how they combine, then the whole system. Every mechanism uses the same
shape — what goes in, what happens, what comes out — so the document can be
read one box at a time.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
FIG = HERE / "figures"
F = json.loads((HERE / "facts.json").read_text())

NAVY = RGBColor(0x1B, 0x2C, 0x47)
GREY = RGBColor(0x5F, 0x63, 0x68)
RED = RGBColor(0xB0, 0x3A, 0x3A)


# ------------------------------------------------------------- primitives
def footer_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = GREY


def mono(par, text, size=9, color=None):
    r = par.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    r.font.color.rgb = color or NAVY
    return r


def para(doc, text="", size=10.5, bold=False, italic=False, color=None,
         space_after=6, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def rich(doc, parts, size=10.5, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for text, style in parts:
        if style == "code":
            mono(p, text, size=size - 1.3)
            continue
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = style == "b"
        r.italic = style == "i"
    return p


def bullets(doc, items, size=10.5, numbered=False):
    for it in items:
        p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        p.paragraph_format.space_after = Pt(2)
        if isinstance(it, tuple):
            head, rest = it
            r = p.add_run(head); r.bold = True; r.font.size = Pt(size)
            r2 = p.add_run(rest); r2.font.size = Pt(size)
        else:
            r = p.add_run(it); r.font.size = Pt(size)


def code(doc, text, size=8.8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Mm(6)
    mono(p, text, size=size)
    return p


def table(doc, headers, rows, widths=None, size=9.2):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(size)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            if val.startswith("`") and val.endswith("`"):
                mono(p, val.strip("`"), size=size - 0.6)
            else:
                r = p.add_run(val); r.font.size = Pt(size)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Mm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return t


def figure(doc, name, caption, width_mm=162):
    path = FIG / name
    if not path.exists():
        raise SystemExit(f"missing figure {path}")
    doc.add_picture(str(path), width=Mm(width_mm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption); r.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = GREY
    c.paragraph_format.space_after = Pt(12)


def h1(doc, text, page_break=True):
    if page_break:
        doc.add_page_break()
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = NAVY


def h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = NAVY


def piece(doc, name, one_line, detail, where):
    """A Level 1 entry: what a thing IS."""
    h2(doc, name)
    para(doc, one_line, size=11, bold=True, space_after=4)
    para(doc, detail, space_after=4)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Where it lives:  "); r.font.size = Pt(9.5); r.italic = True
    r.font.color.rgb = GREY
    mono(p, where, size=8.8)


def mech(doc, num, title, goes_in, steps, comes_out, file, fn, why, breaks):
    """A Level 2 entry: how a mechanism WORKS. The unit of this document."""
    h2(doc, f"{num}. {title}")
    rows = [("WHAT GOES IN", goes_in)]
    for i, s in enumerate(steps, 1):
        rows.append((f"STEP {i}", s))
    rows.append(("WHAT COMES OUT", comes_out))
    table(doc, ["", "What happens"], rows, widths=[32, 133], size=9.2)

    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    r = p.add_run("The code:  "); r.font.size = Pt(9.5); r.italic = True
    r.font.color.rgb = GREY
    mono(p, file, size=8.8)
    r2 = p.add_run("   →   "); r2.font.size = Pt(9)
    mono(p, fn, size=8.8)

    rich(doc, [("Why it exists. ", "b"), (why, "")], space_after=4)
    rich(doc, [("What breaks without it. ", "b"), (breaks, "")], space_after=12)


def build(out_path):
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Mm(210), Mm(297)
    s.left_margin = s.right_margin = Mm(20)
    s.top_margin = s.bottom_margin = Mm(18)
    footer_page_number(s)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    _cover(doc)
    _how_to_read(doc)
    _level1(doc)
    _level2(doc)
    _level3(doc)
    _level4(doc)
    _unfinished(doc)
    _questions(doc)

    doc.save(out_path)
    return out_path


# --------------------------------------------------------------- sections
def _cover(doc):
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ISU-GeoBot"); r.bold = True; r.font.size = Pt(30)
    r.font.color.rgb = NAVY
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("How It Actually Works"); r.font.size = Pt(15)
    r.font.color.rgb = GREY
    para(doc, "", space_after=14)
    para(doc, "From the smallest pieces upward", size=12, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, space_after=20)
    para(doc,
         "Read this in order. It starts with what a single value is, then how each "
         "mechanism works one at a time, then how those mechanisms combine to answer a "
         "real question, and only at the very end shows the whole system. Nothing here "
         "is described from the design documents — every step was read out of the code "
         "in this repository.",
         size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)


def _how_to_read(doc):
    h1(doc, "How to read this guide")
    para(doc, "Every mechanism in Part 2 uses the same three-part shape:")
    table(doc, ["", "Means"], [
        ("WHAT GOES IN", "the data handed to this piece of code"),
        ("STEP 1, 2, 3...", "what it actually does, in order"),
        ("WHAT COMES OUT", "what the next piece receives"),
    ], widths=[36, 129])
    para(doc,
         "If you can say the IN and the OUT for a mechanism, you understand it well "
         "enough to defend it. The steps are there for when someone asks you to go "
         "deeper.")
    h2(doc, "The four parts")
    bullets(doc, [
        ("Part 1 — The pieces. ", "What a status code, a chunk, an embedding and a "
                                  "feature vector actually are. No code yet."),
        ("Part 2 — The mechanisms. ", "Ten things the system does, one at a time. This "
                                      "is the bulk of the guide."),
        ("Part 3 — Three real questions. ", "Which mechanisms fire, in order, for three "
                                            "different questions."),
        ("Part 4 — The whole system. ", "One diagram, one page. Last, not first."),
    ])
    para(doc,
         "Parts 5 and 6 are what is not finished and the ten questions you should be "
         "able to answer without notes.", space_after=4)


# ------------------------------------------------------------- PART 1
def _level1(doc):
    h1(doc, "PART 1 — The pieces")
    para(doc, "Six things. Once these are clear, everything else is just moving them "
              "around.", italic=True, color=GREY)

    piece(doc, "A status code",
          "One of exactly three strings. Nothing else is allowed.",
          "The system never says where a lecturer is. It says one of three things: "
          "available_consultation, in_scheduled_class, or unavailable_off_schedule. "
          "This is a closed set, checked in code — if anything else appears, the "
          "request fails rather than showing it. The user sees a friendly label "
          "('In Scheduled Class / Lecture'), but the value moving through the system is "
          "one of those three strings.",
          "database/tables-and-structure.sql  (table availability_status)")

    piece(doc, "A schedule block",
          "One row saying: this person, this weekday, this time range, this kind, this "
          "campus.",
          "A lecturer's timetable is stored as rows. Each has a day_of_week, a "
          "start_time and end_time, a block_kind (class, consultation or admin), and "
          "importantly a campus. The campus column is what lets the system notice that "
          "someone is teaching, but not here. There are %s blocks in total."
          % F["schedule_total"],
          "database/tables-and-structure.sql  (table faculty_schedule)")

    piece(doc, "A chunk",
          "A piece of an official document, about 200 tokens long.",
          "Documents like the Student Manual are too long to hand to a language model, "
          "so they are cut into pieces. Each piece is aimed at 200 tokens and may never "
          "exceed 220 — that ceiling is enforced by the database itself. Consecutive "
          "chunks overlap by about 15 percent so a sentence sitting on a boundary is "
          "still findable from either side. There are %s chunks."
          % F["document_chunk"],
          "database/tables-and-structure.sql  (table document_chunk)")

    piece(doc, "An embedding",
          "A list of 384 numbers that represents the meaning of a piece of text.",
          "Every chunk is converted into 384 numbers by a model called all-MiniLM-L6-v2. "
          "Text with similar meaning produces similar numbers. When someone asks a "
          "question, the question is converted the same way, and the system finds the "
          "chunks whose numbers point in the most similar direction. That is the whole "
          "idea behind retrieval — no keyword matching is involved.",
          "machine-learning/ai_api_service.py  (the /embed endpoint)")

    piece(doc, "A feature vector",
          "Eleven numbers describing one moment for one lecturer.",
          "The Random Forest does not see names or rooms. It sees eleven numbers: the "
          "weekday, the time slot, whether this moment falls inside a consultation "
          "block, whether it falls inside a class, whether it is an exam period, whether "
          "a disrupting event falls that day, which phase of the semester it is, a "
          "stable index for the person, and three historical rates — how often they have "
          "been present in this slot, how punctual they are, and how often they leave "
          "early.",
          "machine-learning/feature_engineering.py")

    piece(doc, "A pseudonym",
          "A stand-in identifier used so attendance data is not stored against a name.",
          "The attendance table has no faculty_id column at all. It stores a "
          "pseudonym_id instead, and only one small mapping table can turn that back "
          "into a person. This means attendance can be analysed without the analysis "
          "ever touching a name. Note the correct word is pseudonymised, not anonymised "
          "— the link still exists, so under RA 10173 it is still personal data.",
          "database/tables-and-structure.sql  (tables attendance_record, "
          "faculty_pseudonym_map)")


# ------------------------------------------------------------- PART 2
def _level2(doc):
    h1(doc, "PART 2 — The mechanisms")
    para(doc, "Ten mechanisms, each self-contained. Read one, then stop.",
         italic=True, color=GREY)

    figure(doc, "d1_masking.png",
           "Diagram 1 — Status masking: what survives the boundary and what does not")

    mech(doc, 1, "Status masking — how a prediction becomes a status",
         "The complete reply from the ML service: predicted_class, probabilities, "
         "feature_list, model_version.",
         [
             "Check whether predicted_class is one of the three allowed strings.",
             "If it is not, THROW an error. It never falls back to a default status — "
             "a model producing something unexpected is a model whose output must not "
             "be shown to anyone.",
             "Build a brand-new object containing only three fields: statusCode, "
             "source, and maskedAt.",
             "Purge the original: set predicted_class, probabilities and feature_list "
             "to null, so they cannot travel further by accident.",
             "Hand the probabilities back separately, as their own return value, so "
             "storing them for research has to be a deliberate decision.",
         ],
         "Three fields. No room, no building, no campus, no probability, no feature "
         "vector. That object is what the language model is eventually told about.",
         "backend/src/middleware/privacy-masking-middleware.js", "maskPrediction()",
         "This is the single point where internal detail stops. Everything before it "
         "knows the room and the probabilities; everything after it cannot, because "
         "those fields no longer exist on the object.",
         "The language model would receive the full prediction — including the campus "
         "and the confidence numbers — and would happily mention them in the answer.")

    mech(doc, 2, "The guard override — when the model is skipped entirely",
         "The most recent guard log entry for that lecturer, if there is one today.",
         [
             "Look up the latest guard event: an arrival or a departure.",
             "If the latest event is a DEPARTURE, the person has been observed leaving.",
             "Return unavailable_off_schedule immediately, with source = "
             "'guard_override'.",
             "The Random Forest is never called. There is no prediction to mask.",
         ],
         "The same three-field masked object, but produced without any model involvement.",
         "backend/src/middleware/privacy-masking-middleware.js", "maskOverride()",
         "An observation beats an estimate. If someone watched the lecturer leave, "
         "running a probability model against that would produce worse information, not "
         "better.",
         "The system would predict 'in scheduled class' for someone a guard watched walk "
         "out of the gate ten minutes ago.")

    mech(doc, 3, "The egress filter — catching a location the model invented",
         "The finished sentence the language model wrote, but only when that answer "
         "carried an availability status.",
         [
             "If a course code was included, temporarily blank it out of the test copy, "
             "because a code like 'IT 211' looks exactly like a room number to the next "
             "step.",
             "Test the sentence against two families of pattern: LOCATION patterns "
             "(room 204, 2nd floor, building 3, CCS-301, faculty lounge) and "
             "SPECULATION patterns ('probably in...', 'you could try the...').",
             "If nothing matches, return the answer untouched.",
             "If something matches, THROW THE SENTENCE AWAY and substitute a fixed safe "
             "reply naming only the status.",
             "Record that a hit occurred, so the trigger rate can be measured.",
         ],
         "Either the original answer, or a replacement sentence that gives the status "
         "and explicitly says the system does not disclose physical location.",
         "backend/src/middleware/privacy-masking-middleware.js", "filterEgress()",
         "Masking cleans what goes INTO the model. Nothing cleans what comes OUT. Given "
         "a status and asked 'where can I find her?', a language model will invent a "
         "plausible room number from nothing. This is the only thing standing between "
         "that invention and the user.",
         "The privacy boundary would be breached at the final step, invisibly — the "
         "masking layer would have done its job perfectly and the answer would still "
         "name a room.")
    rich(doc, [
        ("Note the deliberate bluntness. ", "b"),
        ("The patterns are broad on purpose. A false positive costs one templated "
         "sentence; a false negative is a privacy incident. The filter is also applied "
         "ONLY to answers that carried a status — a pure navigation answer is supposed "
         "to name a building, so filtering those would break the map feature.", ""),
    ], space_after=12)

    mech(doc, 4, "The consent gate — deciding whether to answer at all",
         "The lecturer's identifier, once the router has matched a name.",
         [
             "Check the caller is signed in. Anonymous visitors are asked to sign in.",
             "Call the database function faculty_is_answerable(), which requires three "
             "things at once: the lecturer is active, has consented, and has not paused "
             "their own visibility.",
             "If any of those is false, stop here. No schedule is read, no model is "
             "called, nothing is computed.",
             "Return a refusal that says the information is unavailable and offers the "
             "campus map instead.",
         ],
         "Either permission to continue, or a refusal — and no computed estimate "
         "anywhere.",
         "backend/src/services/knowledge-search-service.js",
         "faculty_is_answerable RPC, lines 241-260",
         "Refusing BEFORE computing is different from computing then hiding. A paused "
         "lecturer's estimate is never produced at all, which is what respecting an "
         "objection actually means.",
         "The system would calculate a private estimate for someone who asked not to "
         "participate, and merely decline to display it.")
    rich(doc, [
        ("The wording matters too. ", "b"),
        ("The refusal never says the person opted out. Announcing that a named "
         "individual exercised their right to object discloses a choice they made about "
         "their own data, and could invite exactly the pressure that right exists to "
         "protect them from. Pausing and simply not participating are made "
         "indistinguishable from the outside.", ""),
    ], space_after=12)

    mech(doc, 5, "Name matching — turning 'Prof Santos' into a person",
         "The raw text of the question.",
         [
             "Try an exact match on the full name, among consented, active faculty only.",
             "Try an exact match on any recorded alias, same restriction.",
             "Try a fuzzy match using PostgreSQL trigram similarity, which tolerates "
             "typos and partial names.",
             "If exactly one candidate comes back, use it.",
             "If more than one comes back, STOP the whole pipeline and ask the user "
             "which person they meant.",
         ],
         "Either one lecturer, or a clarifying question. Never a guess between two "
         "people.",
         "database/database-functions.sql", "resolve_faculty_candidates()",
         "Guessing between two similarly-named people means answering a question about "
         "the wrong person — which is a privacy failure, not just an accuracy one.",
         "The system would silently pick whichever name scored fractionally higher and "
         "report a stranger's availability.")
    rich(doc, [
        ("The consent filter is inside the matcher. ", "b"),
        ("An unconsented lecturer cannot even be RESOLVED by name, let alone have a "
         "status computed. That is why the gate in mechanism 4 is a second line of "
         "defence rather than the only one.", ""),
    ], space_after=12)

    mech(doc, 6, "Chunking — cutting a document into retrievable pieces",
         "The full text of an official document, converted to Markdown.",
         [
             "Split into natural blocks — paragraphs, list items, headed sections.",
             "Any single block that is already over the target gets split further.",
             "Pack blocks together until adding one more would exceed 200 tokens.",
             "Measure the size with the EMBEDDING MODEL'S OWN tokenizer, not a "
             "word count, because only the model's count matters.",
             "Carry roughly 15 percent of the previous chunk into the next one as "
             "overlap.",
             "Guarantee every chunk is at or under 220 tokens — the database rejects "
             "anything larger.",
         ],
         "A list of chunks, each with its verified token count, ready to be embedded.",
         "machine-learning/document_knowledge_importer.py", "chunk_document()",
         "The embedding model silently truncates anything past its limit. A chunk over "
         "the ceiling would have a tail that is not represented in its own embedding, so "
         "that text becomes permanently unfindable.",
         "Retrieval would fail for content near the end of long chunks, and nothing "
         "would report an error — the text would simply never come back.")

    mech(doc, 7, "Retrieval — finding the right chunks for a question",
         "The user's question, as text.",
         [
             "Send the question to the ML service to be embedded — the SAME service and "
             "SAME model that embedded the documents.",
             "Pass those 384 numbers to the database function match_document_chunks().",
             "The database compares the question vector against every chunk vector using "
             "exact cosine similarity. There is no approximate index.",
             "Discard anything scoring below 0.25.",
             "Return the top 5 that remain.",
         ],
         "Up to five chunks of document text, each with its similarity score and source "
         "document.",
         "backend/src/services/knowledge-search-service.js", "retrieve()",
         "Using one embedder everywhere is the whole reason this works. If documents "
         "were embedded by one model and questions by another, the two sets of numbers "
         "would live in different spaces and retrieval would return nonsense.",
         "With a mismatched embedder, retrieval degrades invisibly — no crash, just "
         "quietly irrelevant results.")
    rich(doc, [
        ("Why exact search and not an index. ", "b"),
        ("With %s chunks, comparing against all of them is fast. An approximate index "
         "would add a 'did it miss anything?' question to every result for no measurable "
         "speed gain at this size." % F["document_chunk"], ""),
    ], space_after=12)

    mech(doc, 8, "The schedule lookup — and the other-campus case",
         "A lecturer, a moment in time, and the campus the question is about.",
         [
             "Check whether that date has a disrupting institutional event — a holiday "
             "or an examination period. If so, return unavailable_off_schedule "
             "immediately.",
             "Find a schedule block covering that weekday and time for that lecturer.",
             "If there is no block at all, return unavailable_off_schedule.",
             "If there is a block, compare its campus against the campus asked about.",
             "Same campus, and it is a class: return in_scheduled_class.",
             "DIFFERENT campus: return unavailable_off_schedule, and set matched_block "
             "to class_other_campus.",
             "A consultation or admin block returns available_consultation.",
         ],
         "A status code, the kind of block that matched, and whether the day was an "
         "event day.",
         "database/database-functions.sql", "schedule_lookup_status()",
         "A lecturer teaching on the Santiago campus is genuinely unavailable to a "
         "student standing in Echague. Without the campus comparison the system would "
         "confidently say they are in class here.",
         "Students would go looking for someone who is fifty kilometres away.")
    rich(doc, [
        ("What the user sees in that case. ", "b"),
        ("The status code stays unavailable_off_schedule — no fourth state is invented. "
         "Only the LABEL changes, to 'Teaching this period; not scheduled on this "
         "campus'. That tells the student waiting is pointless, without naming a campus, "
         "building or room.", ""),
    ], space_after=6)
    rich(doc, [
        ("And the rule that catches people out. ", "b"),
        ("An empty slot in the timetable is NOT availability. Not teaching is not the "
         "same as free — the lecturer may be in a meeting, doing research, or off "
         "campus. Only a declared consultation block, or attendance evidence, produces "
         "an 'available' answer.", ""),
    ], space_after=12)

    mech(doc, 9, "Timezone conversion — the one place time is translated",
         "A timestamp of the current moment, in UTC.",
         [
             "Convert it once, using the configured campus timezone (Asia/Manila).",
             "Produce a 'naive local' value — the wall-clock time on campus, with no "
             "timezone attached.",
             "Every downstream feature is derived from THAT value: the weekday, the time "
             "slot, the calendar date, the semester phase.",
         ],
         "A campus-local timestamp that both the schedule lookup and the model agree on.",
         "backend/src/services/faculty-presence-service.js", "toCampusLocalNaive()",
         "The model was trained on campus-local times. If serving sends anything else, "
         "the model is asked about a different moment than the one the user means.",
         "This exact bug happened: the code sent UTC, so 07:00 Monday in Echague arrived "
         "at the model as 23:00 SUNDAY. Not just the wrong hour — the wrong day.")

    mech(doc, 10, "The location tag — how the map moves safely",
         "The sentence the language model produced, plus the list of real campus "
         "locations.",
         [
             "Look for a marker of the form [LOCATION: some-slug] at the end.",
             "If there is none, the map does not move.",
             "Remove the marker from the text, so the user never sees it.",
             "Look the slug up in the real location list.",
             "If the slug is not a real location, log a warning and IGNORE it.",
             "Only if it matches a real record does the map pan.",
         ],
         "The cleaned answer, and either a verified location or nothing.",
         "backend/src/services/knowledge-search-service.js", "extractLocationTag()",
         "The language model is allowed to REQUEST that the map moves, but never to "
         "decide where. The server checks the request against real data first, so the "
         "model cannot invent a place.",
         "A model that hallucinated a location id could send the map anywhere, or the "
         "raw tag could leak into the visible answer.")


# ------------------------------------------------------------- PART 3
def _level3(doc):
    h1(doc, "PART 3 — Three real questions")
    para(doc, "Now put the mechanisms together. Notice what is the same across all "
              "three, and what is different.", italic=True, color=GREY)

    figure(doc, "d3_three_questions.png",
           "Diagram 3 — The same pipeline, three different questions")

    h2(doc, 'Question A — "Where is the library?"')
    bullets(doc, [
        "The router sees a place word and no faculty name, so the category is "
        "campus_navigation.",
        "Mechanism 7 runs: the question is embedded and the top chunks come back — one "
        "of them is the generated place-card for the University Library.",
        "The campus location list loads alongside it.",
        "No consent gate, no schedule, no model — nothing about a person is involved.",
        "The prompt is built from the chunks plus the location list.",
        "The model answers and ends with [LOCATION: university-library].",
        "Mechanism 10 validates that slug and the map pans to the library.",
        "The egress filter does NOT run, because no status was disclosed — and this "
        "answer is supposed to name a building.",
    ], numbered=True)

    h2(doc, 'Question B — "What is the passing grade?"')
    bullets(doc, [
        "No place word, no faculty name, so the category is general_institutional.",
        "Mechanism 7 runs exactly as before — the same embedding, the same search.",
        "The chunks that come back are from the Student Manual grading section.",
        "Again: no consent gate, no schedule, no model.",
        "The model answers from those chunks: 3.00, equivalent to 75-76 percent.",
        "The map does not move, because no location tag was produced.",
    ], numbered=True)
    rich(doc, [
        ("The point of comparing A and B. ", "b"),
        ("They take the SAME path. Routing did not send one to a document search and the "
         "other somewhere else — retrieval ran for both. The category only changed "
         "whether the map was allowed to move.", ""),
    ], space_after=10)

    h2(doc, 'Question C — "Is Professor X available?"')
    figure(doc, "d2_availability.png",
           "Diagram 2 — The availability decision, in full")
    bullets(doc, [
        "Mechanism 5 matches the name against consented faculty. One match, so continue "
        "— two matches would have stopped here with a question.",
        "Mechanism 4 checks the caller is signed in, then checks consent and visibility. "
        "Both must pass before anything is computed.",
        "Mechanism 2 checks the guard log. If the last entry is a departure, the answer "
        "is decided here and the model is never called.",
        "Mechanism 9 converts the current time to campus-local.",
        "Mechanism 8 looks up the schedule, including the campus comparison.",
        "The attendance history rates are fetched.",
        "The eleven features go to the Random Forest, which returns one of three classes.",
        "Mechanism 1 masks that prediction down to a status code and destroys the rest.",
        "Retrieval ALSO ran, in parallel, exactly as in questions A and B.",
        "The prompt is built from the chunks AND the masked status.",
        "The model writes a sentence.",
        "Mechanism 3 scans that sentence for location leakage and replaces it if found.",
        "The map does NOT move — the question was about a person, not a place.",
    ], numbered=True)
    rich(doc, [
        ("That last point is a real bug that was fixed. ", "b"),
        ("The map used to pan to whichever building happened to be retrieved. Asking "
         "'Is Professor Alado available?' moved the map to a building called 'Alamario' "
         "while the answer said no information was available. The map focus is now "
         "gated on the question category.", ""),
    ], space_after=6)


# ------------------------------------------------------------- PART 4
def _level4(doc):
    h1(doc, "PART 4 — The whole system")
    para(doc, "Only now, once the pieces make sense.", italic=True, color=GREY)
    figure(doc, "d4_whole.png", "Diagram 4 — Everything, on one page")
    para(doc,
         "A React website talks to an Express server, and only to that server. The "
         "server decides what a question needs, then gathers it from three places: a "
         "PostgreSQL database holding people, places and documents; a small Python "
         "service that turns text into numbers and runs the Random Forest; and Groq, "
         "which writes the final sentence. The server assembles the answer, masks "
         "anything private, filters the result, and sends back one clean object.")
    rich(doc, [
        ("The one rule worth memorising: ", "b"),
        ("the browser never talks to the database, the ML service or the language model "
         "directly. That is why no API key and no table name ever reaches the user's "
         "computer.", ""),
    ])
    h2(doc, "Where things live")
    table(doc, ["Part", "Folder", "Runs on"], [
        ("Website", "`frontend/`", "the user's browser"),
        ("Server", "`backend/`", "Node, port 4000"),
        ("Database + functions", "`database/`", "Supabase PostgreSQL"),
        ("ML service and training", "`machine-learning/`", "Python, port 5001"),
    ], widths=[46, 60, 59])


# ------------------------------------------------------------- PART 5
def _unfinished(doc):
    h1(doc, "PART 5 — What is not finished")
    para(doc,
         "Say these before an examiner finds them. Being able to explain a limitation "
         "is much stronger than being surprised by it.", italic=True, color=GREY)
    table(doc, ["What", "The honest answer"], [
        ("The attendance data is invented",
         "All %s attendance records are synthetic. The advisor ruled out using real "
         "Daily Time Records because they are sensitive personal data. So every "
         "availability accuracy figure is a SIMULATION result: it shows the pipeline "
         "recovers patterns that were deliberately put in, not that it predicts real "
         "lecturers." % F["attendance"]),
        ("RAGAS has never produced a score",
         "The answers were generated — %s results across two runs — but scoring them "
         "needs about a million tokens from the judging model and the free plan allows "
         "200,000 a day. It is a quota limit, not a broken pipeline."
         % F["eval_result"]),
        ("Faculty validation has barely started",
         "%s entries. The portal and its permissions work; what is missing is "
         "consented lecturers actually using it." % F["faculty_validation"]),
        ("The consultation blocks were generated, not declared",
         "%s consultation blocks were created from empty gaps in the timetable rather "
         "than hours lecturers published. This is exactly the 'a gap is not "
         "availability' rule being broken in the data, and it makes one input feature "
         "nearly identical to the answer the model is trying to predict."
         % F["schedule_consult"]),
        ("Guard events are test data",
         "%s rows were created in a five-minute window by one person testing the "
         "portal, including entries for synthetic lecturers, yet they are marked as "
         "real." % F["guard_events"]),
        ("No claim of legal compliance",
         "The system is designed in accordance with RA 10173 principles. It is NOT "
         "claimed to be compliant — that is a legal determination needing a formal "
         "impact assessment, which is outside a thesis."),
    ], widths=[52, 113])


# ------------------------------------------------------------- PART 6
def _questions(doc):
    h1(doc, "PART 6 — Ten questions to answer without notes")
    qa = [
        ("What happens when someone asks a question?",
         "It goes to the server, which classifies it, searches the document chunks, and "
         "— only if the question is about a consented lecturer and the user is signed in "
         "— computes a masked availability status. All of that goes into one prompt, the "
         "language model writes the answer, and a filter checks it before it is sent."),
        ("How does status masking work?",
         "The model returns a predicted class, probabilities and the feature list. "
         "Masking checks the class is one of three allowed values, throws if it is not, "
         "builds a new object with only the status code and a timestamp, and sets the "
         "probabilities and features to null. Three fields survive; everything else is "
         "destroyed."),
        ("Why do you need a filter on the output as well?",
         "Because masking cleans the input to the language model, not its output. Given "
         "a status and asked where someone is, the model will invent a room number. The "
         "filter scans for that and replaces the sentence."),
        ("Why is an empty period not shown as free?",
         "Not teaching is not the same as available. They could be in a meeting, doing "
         "research, or off campus. Saying 'free' would be a claim about a person's "
         "whereabouts with no evidence behind it."),
        ("What happens if a lecturer is teaching on another campus?",
         "The status stays unavailable_off_schedule, but the label becomes 'Teaching "
         "this period; not scheduled on this campus'. No new state is invented, and no "
         "campus or building is named."),
        ("What if a lecturer does not want to take part?",
         "Then nothing is computed about them at all — the check happens before the "
         "schedule is even read. They can also pause their own visibility whenever they "
         "want. The refusal deliberately does not say they opted out."),
        ("Why is 220 tokens a hard limit?",
         "Because the embedding model silently truncates longer text. Anything past the "
         "limit would not be represented in its own embedding, so it could never be "
         "found. The database enforces the ceiling."),
        ("Why one Python service instead of doing it in Node?",
         "Because the same embedding model must be used for documents, questions and "
         "evaluation. If two different embedders were used, the numbers would not be "
         "comparable and search would quietly return the wrong things."),
        ("Is your accuracy figure real?",
         "It is a simulation result and we label it that way. The attendance is "
         "synthetic because real time records were ruled out on privacy grounds. It "
         "shows the pipeline recovers the patterns that were injected, and that a "
         "schedule-only rule structurally cannot."),
        ("What is your biggest limitation?",
         "That the availability evaluation rests on synthetic attendance and that "
         "validation has barely begun. Both come from the same cause: we cannot collect "
         "real presence data without consent we do not yet have."),
    ]
    for i, (q, a) in enumerate(qa, 1):
        rich(doc, [(f"{i}. {q}", "b")], space_after=2)
        rich(doc, [(a, "")], space_after=9)

    para(doc, "")
    rich(doc, [
        ("If you do not know a number, say you will check it. ", "b"),
        ("Knowing the edge of your own evidence reads far better than filling the gap "
         "with something plausible.", "i"),
    ])


if __name__ == "__main__":
    out = ROOT / "ISU-GeoBot-How-It-Works.docx"
    build(out)
    print(f"written: {out}")
