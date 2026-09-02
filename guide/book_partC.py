"""Book, Parts VI-VIII: frontend, running it, defense. Plus the assembler."""

from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from book_partA import filecard, h3, part1, part2
from book_partB import part3, part4, part5
from build_guide import (F, GREY, NAVY, bullets, code, figure,
                         footer_page_number, h1, h2, mono, para, rich, table)


# =====================================================================
#  PART VI — THE FRONTEND
# =====================================================================
def part6(doc):
    h1(doc, "PART VI — The frontend")
    para(doc, "What the user actually touches.", italic=True, color=GREY)

    h1(doc, "Chapter 17 — Structure and build")

    h2(doc, "17.1 How it boots")
    para(doc,
         "`index.html` provides a single div. `src/main.jsx` mounts React into it. "
         "`src/App.jsx` sets up the router and the theme context. Vite serves it in "
         "development and bundles it for production.")

    h2(doc, "17.2 The one rule the frontend follows")
    para(doc,
         "Every network call goes through `src/frontend-utilities/backendApiClient.js`. "
         "No component builds its own fetch. That wrapper adds the base URL, attaches "
         "the bearer token when signed in, and normalises errors.")
    rich(doc, [
        ("Why this matters more than it looks. ", "b"),
        ("It is the reason no component can accidentally call Supabase directly and "
         "bypass the server. `supabaseClient.js` exists, but it is used only for "
         "sign-in — never to read data.", ""),
    ])

    h2(doc, "17.3 Styling")
    para(doc,
         "Tailwind utility classes, with design tokens in `src/styles/tokens.css` for "
         "colour and spacing. Dark mode is a context provider in `themeContext.jsx`. "
         "There is no component library — the primitives in `components/ui-primitives/` "
         "are hand-written and small.")

    h1(doc, "Chapter 18 — The five surfaces")

    h2(doc, "18.1 The main assistant")
    para(doc, "`components/main-assistant/` — the primary experience.")
    filecard(doc, "MainAssistantWorkspace.jsx",
             "the shell that holds chat and map state together, so an answer can move "
             "the map")
    filecard(doc, "ChatbotMessagePanel.jsx",
             "renders the conversation, the status badge and the sources")
    filecard(doc, "InteractiveCampusMap.jsx",
             "the Leaflet map: markers, pan-to-location, and the walking route overlay")
    filecard(doc, "PlacesDirectoryDrawer.jsx · PlaceDetailCard.jsx",
             "browse locations without asking a question")
    filecard(doc, "FloatingChatDock.jsx", "the compact launcher used on the landing page")
    para(doc,
         "`campusRoutingService.js` and `campusRoadGraph.json` compute a walking path "
         "across the campus path network — this runs entirely in the browser and does "
         "not involve the server.")

    h2(doc, "18.2 The landing page")
    para(doc,
         "`components/landing-page/` — fifteen components that explain the system to a "
         "visitor, including diagrams of the RAG pipeline and the masking protocol. This "
         "is presentation only; it makes no privileged API calls.")

    h2(doc, "18.3 The admin portal")
    filecard(doc, "AdminDashboard.jsx", "the shell and navigation")
    filecard(doc, "CampusMapEditor.jsx",
             "click the map to place or move a location")
    filecard(doc, "CampusLocationManager.jsx",
             "the list view: create, edit, unpublish, republish, delete, reindex")
    filecard(doc, "AdminFacultyValidationPanel.jsx", "review validation entries")
    filecard(doc, "AdminSecurityPresencePanel.jsx", "review guard activity")
    para(doc,
         "Every write here calls an admin route which requires the admin or researcher "
         "role and writes a `poi_audit` row. Hiding these screens is not the security "
         "boundary — the server is.")

    h2(doc, "18.4 The validator portal")
    para(doc, "`components/faculty-validation-portal/` — for consented lecturers.")
    filecard(doc, "FacultyAnswerValidationList.jsx",
             "shows what the system estimated and asks whether it was right: correct, "
             "partially correct, or incorrect")
    filecard(doc, "FacultyPrivacyToggleCard.jsx",
             "pause or resume your own availability, without asking the researchers")
    para(doc,
         "This portal is the only source of real-world ground truth for the availability "
         "arm. It needs no attendance data — a lecturer simply says whether the estimate "
         "matched reality.")

    h2(doc, "18.5 The guard portal")
    filecard(doc, "SecurityGuardAttendanceDashboard.jsx",
             "log an arrival or a departure against a consented lecturer")
    para(doc,
         "A departure is the one input that overrides the model entirely — see 14.3.")


# =====================================================================
#  PART VII — RUNNING AND PROVING IT
# =====================================================================
def part7(doc):
    h1(doc, "PART VII — Running and proving it")

    h1(doc, "Chapter 19 — Configuration")
    para(doc,
         "Every setting is an environment variable, read once into one object in "
         "`backend/src/utilities/configuration.js`. Nothing is hardcoded.")
    table(doc, ["Variable", "Default", "What it controls"], [
        ("`DEMO_MODE`", "false", "swap all external services for mocks"),
        ("`PORT`", "4000", "the API port"),
        ("`CORS_ORIGINS`", "localhost:5173", "which origins may call the API"),
        ("`SUPABASE_URL`", "required", "the database"),
        ("`SUPABASE_SERVICE_ROLE_KEY`", "required", "server-only key, bypasses RLS"),
        ("`SUPABASE_ANON_KEY`", "—", "used only to verify user tokens"),
        ("`DB_SCHEMA`", "geobot", "the schema name"),
        ("`ML_BASE_URL`", "127.0.0.1:5001", "loopback — the ML service is not public"),
        ("`GROQ_API_KEY`", "required", "the language model"),
        ("`GROQ_MODEL`", "pinned", "recorded on every evaluation run"),
        ("`RETRIEVAL_TOP_K`", "5", "how many chunks to retrieve"),
        ("`RETRIEVAL_SIMILARITY_FLOOR`", "0.25", "minimum similarity to keep"),
        ("`CAMPUS_TIMEZONE`", "Asia/Manila", "the one conversion point"),
        ("`CAMPUS`", "echague", "which campus this deployment answers for"),
        ("`RATE_CHAT_MAX`", "15", "chat requests per minute — a privacy control"),
    ], widths=[58, 32, 75])
    rich(doc, [
        ("`CAMPUS` deserves a sentence. ", "b"),
        ("Eleven CCSICT lecturers also teach in Santiago City. \"Available\" is always "
         "relative to a place, so the system has to know where \"here\" is. A class on "
         "the other campus is neither availability nor absence.", ""),
    ])

    h1(doc, "Chapter 20 — Running it")
    para(doc, "Three processes, in this order:")
    code(doc,
         "cd machine-learning && python ai_api_service.py      # port 5001\n"
         "cd backend          && npm start                     # port 4000\n"
         "cd frontend         && npm run dev                   # port 5173")
    para(doc, "`start.bat` and `stop.bat` at the repository root do all three.")
    h2(doc, "20.1 Preflight")
    code(doc, "cd backend && npm run preflight")
    para(doc,
         "Checks the Supabase URL and key, that the schema is reachable, how many "
         "locations are published, coordinate provenance, how many chunks are embedded, "
         "that row-level security is forced on every table, whether the ML service is up "
         "and which embedder it has loaded, and whether the Groq key is set. Run it "
         "before any demo.")
    h2(doc, "20.2 First-time database setup")
    code(doc,
         "python database/apply-database-migrations.py            # dry run\n"
         "python database/apply-database-migrations.py --initial --run   # FIRST install\n"
         "python database/apply-database-migrations.py --run      # migrations only")
    rich(doc, [
        ("The `--initial` flag is dangerous on purpose. ", "b"),
        ("`tables-and-structure.sql` opens by DROPPING the whole schema. That is correct "
         "for a first install and catastrophic for the second, so the script refuses "
         "`--initial` if the database already holds locations.", ""),
    ])

    h1(doc, "Chapter 21 — Testing")
    table(doc, ["Suite", "Covers"], [
        ("`campus-location-protocol.test.js`",
         "the [LOCATION: id] tag, validation, and when the map may move"),
        ("`chat-conversation.test.js`",
         "history sanitisation and prompt assembly"),
        ("`privacy-masking-security.test.js`",
         "the five masking properties and the egress filter"),
        ("`user-authorization.test.js`",
         "role gates, the audit trail, and prompt-injection attempts"),
        ("`temporal-consistency.test.js`",
         "timezone boundaries and semester phase, cross-checked against Python"),
        ("`live-stack.test.js`",
         "the real Supabase, ML service and Groq"),
    ], widths=[58, 107])
    code(doc,
         "cd backend && npm test                  # fast suite, DEMO_MODE\n"
         "cd backend && LIVE_TESTS=1 npm test     # + the live stack\n"
         "python machine-learning/test_schedule_rule_contract.py")
    h2(doc, "21.1 Why the live suite exists")
    para(doc,
         "Every other test file opens with `DEMO_MODE = true`. Three real integration "
         "bugs survived a fully passing suite because the demo retriever never returned "
         "the chunk that triggered them: the map panned on a person question, the "
         "history features were silently zero, and punch timestamps were eight hours "
         "out. None are exotic — they are ordinary integration bugs, invisible to a "
         "suite that never integrates.")

    h1(doc, "Chapter 22 — Evaluation")
    h2(doc, "22.1 The design")
    para(doc,
         "A pre-registered test set of %s questions with ground-truth answers written "
         "from the source documents. Each question is run through BOTH arms, back to "
         "back, so upstream slowness affects both equally. That produces %s "
         "`eval_result` rows across two runs." % (F["eval_query"], F["eval_result"]))
    h2(doc, "22.2 The gates")
    bullets(doc, [
        "The harness refuses to run in DEMO_MODE.",
        "It refuses if the judge model equals the generator — self-evaluation.",
        "It now refuses a judge that does not exist on the account, before spending a "
        "whole run.",
        "It refuses any test set containing availability questions while attendance is "
        "synthetic.",
        "`004_eval_queries.sql` refuses to register a half-written test set, because "
        "choosing the category mix after seeing results is p-hacking.",
    ])
    h2(doc, "22.3 The metrics")
    table(doc, ["Metric", "Asks"], [
        ("Context Precision", "were the retrieved chunks relevant?"),
        ("Context Recall", "did retrieval find what the ground truth needed?"),
        ("Faithfulness", "is the answer supported by the retrieved context?"),
        ("Answer Relevancy", "does the answer address the question?"),
        ("Response Time", "median and p95, with a per-stage breakdown"),
    ], widths=[46, 119])
    rich(doc, [
        ("An honest note on Context Precision. ", "b"),
        ("Both arms share a retriever, so that metric measures the same retrieval step "
         "twice and is expected to be roughly flat. Report it flat and explain why — a "
         "panelist who knows RAGAS will respect that far more than four bars that all "
         "happen to rise.", ""),
    ])


# =====================================================================
#  PART VIII — DEFENSE
# =====================================================================
def part8(doc):
    h1(doc, "PART VIII — Defending it")

    h1(doc, "Chapter 23 — Three questions, traced end to end")
    figure(doc, "d3_three_questions.png",
           "The same pipeline, three different questions")

    h2(doc, '23.1 "Where is the university library?"')
    bullets(doc, [
        "Router: navigation words, no faculty name → `campus_navigation`.",
        "The question is embedded and matched against the corpus; the library's "
        "place-card comes back among the top five.",
        "The campus location list loads in parallel.",
        "No consent gate, no schedule, no model — no person is involved.",
        "The prompt is built from the chunks plus the location list.",
        "The model answers and ends with `[LOCATION: university-library]`.",
        "The slug is validated against real records; the map pans.",
        "The egress filter does NOT run — no status was disclosed, and this answer is "
        "supposed to name a building.",
    ], numbered=True)

    h2(doc, '23.2 "What is the passing grade?"')
    bullets(doc, [
        "Router: no navigation words, no faculty name → `general_institutional`.",
        "The question is embedded and matched — the same code path as 23.1.",
        "Chunks from the Student Manual grading section come back.",
        "No consent gate, no schedule, no model.",
        "The model answers from those chunks: 3.00, equivalent to 75 to 76 percent.",
        "No location tag, so the map does not move.",
    ], numbered=True)
    rich(doc, [
        ("The point of putting these side by side. ", "b"),
        ("They take the SAME path. Routing did not send one to a document search and the "
         "other to a database. Retrieval ran for both. The category only changed whether "
         "the map was allowed to move.", ""),
    ])

    h2(doc, '23.3 "Is Professor Santos available?"')
    figure(doc, "d2_availability.png", "The availability decision in full")
    bullets(doc, [
        "The name is matched against the CONSENTED gazetteer. One match, so continue — "
        "two would have stopped here with a clarifying question.",
        "Gate A: is the caller signed in? If not, decline and offer navigation.",
        "Gate B: `faculty_is_answerable()` — active, consented, not paused.",
        "The guard log is checked. A departure ends it here; the model is never called.",
        "The current time is converted to campus-local, once.",
        "The timetable is read, including the campus comparison.",
        "The historical attendance rates are fetched.",
        "Eleven features go to the Random Forest; three classes come back.",
        "Masking projects that to one status code and destroys the rest.",
        "Retrieval ALSO ran, in parallel, exactly as in 23.1 and 23.2.",
        "The prompt is built from the chunks AND the masked status.",
        "The model writes a sentence.",
        "The egress filter scans it for invented locations.",
        "The map does NOT move — the question was about a person.",
    ], numbered=True)

    h1(doc, "Chapter 24 — What is not finished")
    para(doc,
         "State these before an examiner finds them. Explaining a limitation is far "
         "stronger than being surprised by one.", italic=True, color=GREY)
    table(doc, ["What", "The honest position"], [
        ("The attendance data is synthetic",
         "All %s records are generated. The advisor ruled out real Daily Time Records "
         "because they are sensitive personal data. Every availability accuracy figure "
         "is therefore a SIMULATION result: it shows the pipeline recovers patterns that "
         "were deliberately injected, not that it predicts real lecturers."
         % F["attendance"]),
        ("RAGAS has produced no score yet",
         "Generation is complete — %s results across two runs — but scoring needs about "
         "1.07 million judge tokens and the free plan allows 200,000 per day. A quota "
         "limit, not a broken pipeline." % F["eval_result"]),
        ("Faculty validation has barely started",
         "%s entries. The portal and its permissions work; what is missing is consented "
         "lecturers using it." % F["faculty_validation"]),
        ("Consultation blocks were generated, not declared",
         "%s consultation blocks were created from empty gaps in the timetable rather "
         "than hours lecturers published. That breaks the 'a gap is not availability' "
         "rule inside the data itself, and makes `is_consultation_hour` nearly identical "
         "to the label the model predicts." % F["schedule_consult"]),
        ("Guard events are test data",
         "%s rows created in a five-minute window by one logger, including entries for "
         "synthetic lecturers, yet marked as real." % F["guard_events"]),
        ("The model may have circular labels",
         "If `label_source` is `schedule_derived`, the evaluation harness itself logs "
         "that such accuracy is 'not evidence that ML beats rule-based lookup'. Check "
         "which artifact is loaded before quoting a number."),
        ("No legal compliance claim",
         "Designed in accordance with RA 10173 principles. NOT claimed compliant — that "
         "needs a formal impact assessment, outside thesis scope."),
    ], widths=[52, 113])

    h1(doc, "Chapter 25 — Questions you should answer without notes")
    qa = [
        ("What happens after someone asks a question?",
         "It reaches POST /api/chat. The server classifies it, searches the document "
         "chunks, and — only if it names a consented lecturer and the caller is signed "
         "in — computes a masked availability status. All of that goes into one prompt, "
         "the language model writes the answer, and a filter checks it before it is "
         "returned."),
        ("What is Enhanced about your RAG?",
         "A Random Forest sits inside the retrieval pipeline. The enhanced arm adds a "
         "live, privacy-masked availability status as a third information source "
         "alongside the retrieved chunks and the campus gazetteer. That is the only "
         "difference between the two arms."),
        ("How does status masking work?",
         "The model returns a predicted class, a probability vector and the feature "
         "list. Masking checks the class against a closed three-value allowlist, throws "
         "if it is not a member, builds a new object with only the status code and a "
         "timestamp, and nulls the probabilities and features on the original. Three "
         "fields survive."),
        ("Why do you need a filter on the output as well as the input?",
         "Masking cleans what the model is told. Nothing constrains what it says. Given "
         "a status and asked where someone is, it will invent a room number. The filter "
         "scans for that and substitutes a safe reply."),
        ("Why is an empty period not shown as free?",
         "Not teaching is not the same as available — they may be in a meeting, doing "
         "research, or off campus. Saying free would be a claim about a person's "
         "whereabouts with no evidence behind it."),
        ("What happens if a lecturer teaches on another campus?",
         "The status stays unavailable_off_schedule; only the label changes, to "
         "'Teaching this period; not scheduled on this campus'. No new state is "
         "invented, and no campus or building is named."),
        ("What if a lecturer does not want to take part?",
         "Nothing is computed about them at all — the check happens before the schedule "
         "is even read. They can pause their own visibility at any time, and the refusal "
         "never says they opted out."),
        ("Why is 220 tokens a hard limit?",
         "Because the embedding model silently truncates beyond it. Text past the limit "
         "would not be represented in its own embedding and could never be retrieved. "
         "The database enforces the ceiling."),
        ("Why one Python service rather than doing it in Node?",
         "The same embedder must be used for documents, questions and evaluation. Two "
         "embedders produce two vector spaces and the comparison becomes meaningless. "
         "Also, scikit-learn and sentence-transformers are Python libraries; Node cannot "
         "load a joblib artifact."),
        ("Why pgvector rather than a dedicated vector database?",
         "The vectors live in the same PostgreSQL instance as the relational data, so a "
         "chunk can carry a foreign key to the location it describes and one backup "
         "covers everything. At %s chunks, exact cosine search is fast and removes any "
         "approximate-recall question." % F["document_chunk"]),
        ("How do you know retrieval works?",
         "A pre-registered set of %s questions with ground truth written from the source "
         "documents, run through both arms. Generation is complete; scoring is pending "
         "on an API quota." % F["eval_query"]),
        ("Is your accuracy figure real?",
         "It is a simulation result and we label it that way. Real time records were "
         "ruled out on privacy grounds. It shows the pipeline recovers injected "
         "patterns, and that a schedule-only rule structurally cannot."),
        ("What is your biggest limitation?",
         "The availability evaluation rests on synthetic attendance and validation has "
         "barely begun. Both come from one cause: we cannot collect real presence data "
         "without consent we do not yet have."),
        ("What would you do differently?",
         "Secure the consented pilot cohort first. Almost every open item is blocked on "
         "that same dependency, and it is the one thing engineering cannot substitute "
         "for."),
    ]
    for i, (q, a) in enumerate(qa, 1):
        rich(doc, [(f"{i}. {q}", "b")], space_after=2)
        rich(doc, [(a, "")], space_after=9)
    para(doc, "")
    rich(doc, [
        ("If you do not know a number, say you will check it. ", "b"),
        ("Knowing the edge of your own evidence reads far better than filling the gap "
         "with something plausible.", "i"),
    ])

    # ------------------------------------------------------ appendices
    h1(doc, "Appendix A — Where to find things")
    table(doc, ["If you want to understand...", "Open"], [
        ("the whole request", "`backend/src/services/knowledge-search-service.js`"),
        ("how a question is classified", "`backend/src/services/intent-query-router.js`"),
        ("availability", "`backend/src/services/faculty-presence-service.js`"),
        ("privacy", "`backend/src/middleware/privacy-masking-middleware.js`"),
        ("the prompt", "`backend/src/utilities/ai-prompt-templates.js`"),
        ("every setting", "`backend/src/utilities/configuration.js`"),
        ("the schema", "`database/tables-and-structure.sql`"),
        ("the SQL logic", "`database/database-functions.sql`"),
        ("row-level security", "`database/security-and-permissions.sql`"),
        ("the model's features", "`machine-learning/feature_engineering.py`"),
        ("how the model is trained", "`machine-learning/train_availability_model.py`"),
        ("how documents are ingested", "`machine-learning/document_knowledge_importer.py`"),
        ("the ML service endpoints", "`machine-learning/ai_api_service.py`"),
        ("the evaluation harness", "`backend/src/services/evaluation-runner.js`"),
    ], widths=[68, 97])

    h1(doc, "Appendix B — Glossary")
    table(doc, ["Term", "Meaning here"], [
        ("Enhanced RAG", "retrieval plus a Random Forest availability status fused into "
                         "the same prompt — the thesis contribution"),
        ("Status Masking Protocol", "the boundary that projects internal state down to "
                                    "one of three status codes"),
        ("Egress filter", "the regex screen applied to the generated answer when a "
                          "status was disclosed"),
        ("Tri-state presence", "confirmed_on_campus, confirmed_off_campus, unknown"),
        ("Context Fusion", "merging retrieved chunks, the gazetteer and the masked "
                           "status into one prompt"),
        ("Gazetteer", "the consented faculty name list and the campus location list"),
        ("Chunk", "a piece of a document, target 200 tokens, ceiling 220"),
        ("Embedding", "384 numbers representing the meaning of a piece of text"),
        ("pgvector", "the PostgreSQL extension that stores and compares embeddings"),
        ("Pseudonymisation", "replacing identity with a surrogate key — still personal "
                             "data under RA 10173"),
        ("data_origin", "a column on every research table: real or synthetic"),
        ("Train/serve skew", "when training and production compute a feature "
                             "differently — the recurring bug family here"),
        ("Circular labels", "when the label is derived from the same source as the "
                            "features, so the model re-derives its own answer"),
        ("RAGAS", "Context Precision, Context Recall, Faithfulness, Answer Relevancy"),
        ("DEMO_MODE", "swaps database, ML and LLM for in-memory mocks"),
        ("Place-card", "generated prose describing a location, chunked and embedded so "
                       "navigation flows through the same retriever"),
    ], widths=[46, 119])


# =====================================================================
#  ASSEMBLY
# =====================================================================
def _cover(doc):
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ISU-GeoBot"); r.bold = True; r.font.size = Pt(34)
    r.font.color.rgb = NAVY
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("The Complete System Book"); r.font.size = Pt(16)
    r.font.color.rgb = GREY
    para(doc, "", space_after=16)
    para(doc, "Every folder, every service, every mechanism", size=12, italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, space_after=24)
    para(doc,
         "Read this front to back. Part I tells you what the system is and names every "
         "file in the repository. Part II covers the data it stores. Parts III to V "
         "explain the backend, the intelligence and the privacy guarantees. Part VI is "
         "the interface, Part VII is how to run and prove it, and Part VIII is how to "
         "defend it.",
         size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    para(doc,
         "Nothing here comes from the design documents. Every path, function name and "
         "number was read out of this repository and the live database while the "
         "document was being built.",
         size=10, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)


def _toc(doc):
    doc.add_page_break()
    h = doc.add_heading("Contents", level=1)
    for run in h.runs:
        run.font.color.rgb = NAVY
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-2" \\h \\z \\u')
    inner = OxmlElement("w:r"); txt = OxmlElement("w:t")
    txt.text = "Select all (Ctrl+A) then press F9 to build the contents."
    inner.append(txt); fld.append(inner)
    p._p.append(fld)
    para(doc, "")
    para(doc, "In Word: Ctrl+A then F9 fills in this table and every page number.",
         size=9, italic=True, color=GREY)
    h2(doc, "The eight parts")
    table(doc, ["Part", "Chapters", "What it covers"], [
        ("I — Orientation", "1-3", "what it is, the tech stack, every folder"),
        ("II — Data", "4-6", "the schema, the SQL functions, what is stored"),
        ("III — Backend", "7-10", "startup, routes, services, the request pipeline"),
        ("IV — Intelligence", "11-13", "RAG, machine learning, the language model"),
        ("V — Guarantees", "14-16", "privacy, security, time and the calendar"),
        ("VI — Frontend", "17-18", "structure and the five surfaces"),
        ("VII — Running it", "19-22", "configuration, deployment, testing, evaluation"),
        ("VIII — Defense", "23-25", "walkthroughs, limitations, questions"),
    ], widths=[42, 26, 97])


def build_book(out_path):
    doc = Document()
    s = doc.sections[0]
    s.page_width, s.page_height = Mm(210), Mm(297)
    s.left_margin = s.right_margin = Mm(20)
    s.top_margin = s.bottom_margin = Mm(18)
    footer_page_number(s)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    _cover(doc)
    _toc(doc)
    part1(doc)
    part2(doc)
    part3(doc)
    part4(doc)
    part5(doc)
    part6(doc)
    part7(doc)
    part8(doc)

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    from pathlib import Path
    out = Path(__file__).resolve().parent.parent / "ISU-GeoBot-The-Complete-Book.docx"
    build_book(out)
    print(f"written: {out}")
