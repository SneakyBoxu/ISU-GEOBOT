"""Book, Parts I-III: orientation, technology, repository, data, backend."""

from __future__ import annotations

from build_guide import (F, GREY, NAVY, bullets, code, figure, h1, h2, mech,
                         mono, para, rich, table)
from docx.shared import Pt


def h3(doc, text):
    h = doc.add_heading(text, level=3)
    for r in h.runs:
        r.font.color.rgb = NAVY


def filecard(doc, path, what):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Pt(12)
    mono(p, path, size=8.6)
    r = p.add_run("  —  " + what)
    r.font.size = Pt(9.5)


# =====================================================================
#  PART I — ORIENTATION
# =====================================================================
def part1(doc):
    h1(doc, "PART I — Orientation")
    para(doc, "What the system is, what it is built from, and where every file "
              "lives.", italic=True, color=GREY)

    # ---------------------------------------------------------- ch 1
    h1(doc, "Chapter 1 — What ISU-GeoBot is")

    h2(doc, "1.1 The problem")
    para(doc,
         "A student on the ISU Echague campus has three kinds of question that are "
         "annoying to answer today. Where is a building they have never been to? What "
         "does a university rule actually say — the passing grade, the attendance rule, "
         "how to drop a subject? And is a particular lecturer likely to be free right "
         "now, or is walking across campus a wasted trip?")
    para(doc,
         "The first is answered by asking someone. The second is answered by finding a "
         "PDF nobody has opened. The third is answered by walking there and hoping. "
         "ISU-GeoBot answers all three from one chat box, with a campus map beside it.")

    h2(doc, "1.2 The three question types")
    table(doc, ["Question", "Example", "What answers it"], [
        ("Where is a place?", "\"Where is the university library?\"",
         "document search over generated place descriptions, plus the map"),
        ("What does a policy say?", "\"What is the passing grade?\"",
         "document search over the Student Manual and School Calendar"),
        ("Is a lecturer free?", "\"Is Professor Santos available?\"",
         "timetable + calendar + attendance history + a Random Forest"),
    ], widths=[38, 58, 69])

    h2(doc, "1.3 The shape of the whole thing, once")
    para(doc,
         "Here is the entire system in one picture. Do not try to absorb it — it is here "
         "so you have somewhere to put the detail that follows. Chapter 3 names every "
         "file inside these boxes.")
    figure(doc, "d4_whole.png", "The whole system, before any of the detail")

    h2(doc, "1.4 The research contribution")
    rich(doc, [
        ("The thesis contribution is called ", ""), ("Enhanced RAG", "b"),
        (". Ordinary RAG — Retrieval-Augmented Generation — searches documents and "
         "hands the results to a language model. Enhanced RAG adds a second, live "
         "source: a Random Forest classifier that estimates a lecturer's availability, "
         "whose output is privacy-masked and fused into the same prompt as the retrieved "
         "text.", ""),
    ])
    para(doc,
         "So the system is evaluated as two arms. The STANDARD arm is ordinary RAG. The "
         "ENHANCED arm is identical in every respect except that an availability block "
         "is appended to the prompt. That single difference is the entire experiment, "
         "and keeping it to a single difference is why the comparison is meaningful.")

    h2(doc, "1.5 The principle everything is built around")
    para(doc, "A status, never a location.", size=15, bold=True, color=NAVY,
         space_after=8)
    para(doc,
         "The system may say a lecturer is estimated to be in a scheduled class. It must "
         "never say which room, which building, which campus, or where they were an hour "
         "ago. This is not a preference — a named person's real-time whereabouts is "
         "personal information under RA 10173, the Philippine Data Privacy Act, and a "
         "sequence of present-moment answers is a movement history.")
    para(doc,
         "That principle is enforced in code at four separate places, described in "
         "Chapter 14. It is the thing to say first if an examiner asks about privacy.")

    h2(doc, "1.6 Who uses it")
    table(doc, ["Role", "What they can do", "Where"], [
        ("Anonymous visitor", "map, navigation and policy questions",
         "the main assistant"),
        ("Signed-in student", "all of the above, plus availability questions",
         "the main assistant"),
        ("Faculty / validator", "confirm whether an estimate was right; pause own "
                                "visibility", "the validator portal"),
        ("Security guard", "log arrivals and departures", "the guard portal"),
        ("Admin / researcher", "edit campus locations, view audit trail",
         "the admin portal"),
    ], widths=[34, 78, 53])

    # ---------------------------------------------------------- ch 2
    h1(doc, "Chapter 2 — The technology stack")
    para(doc, "Every library the system depends on, what it does, and why it is "
              "there.", italic=True, color=GREY)

    h2(doc, "2.1 The four runtimes")
    table(doc, ["Runtime", "Version", "Responsibility"], [
        ("Browser (JavaScript)", "React 18.3", "the interface and the map"),
        ("Node.js", "20", "the API server — all orchestration and all secrets"),
        ("PostgreSQL", "Supabase-hosted", "relational data AND vector search"),
        ("Python", "3.11", "embeddings, the Random Forest, ingestion, evaluation"),
    ], widths=[42, 34, 89])
    para(doc,
         "There is a fifth participant that is not ours: Groq, an external API that "
         "hosts the language model. Nothing else leaves the campus network.")

    h2(doc, "2.2 Frontend dependencies")
    table(doc, ["Package", "Version", "What it does here"], [
        ("`react` / `react-dom`", "18.3.1", "the component model and rendering"),
        ("`react-router-dom`", "6.26", "client-side routing between the portals"),
        ("`leaflet`", "1.9.4", "the interactive map engine"),
        ("`react-leaflet`", "4.2.1", "React bindings for Leaflet"),
        ("`lucide-react`", "0.427", "icon set used across the interface"),
        ("`@supabase/supabase-js`", "2.45", "sign-in only — never data reads"),
        ("`vite`", "5.4", "dev server and production bundler"),
        ("`tailwindcss`", "3.4", "utility-first styling"),
        ("`postcss` / `autoprefixer`", "8.4 / 10.4", "CSS processing for Tailwind"),
    ], widths=[52, 28, 85])
    rich(doc, [
        ("Note what is absent. ", "b"),
        ("There is no state-management library, no component framework, no chart "
         "library beyond a small hand-written bar chart. The frontend is deliberately "
         "plain so that the interesting behaviour is all server-side and therefore "
         "testable.", ""),
    ])

    h2(doc, "2.3 Backend dependencies")
    table(doc, ["Package", "Version", "What it does here"], [
        ("`express`", "4.19", "the HTTP server and routing"),
        ("`@supabase/supabase-js`", "2.45", "database client, service-role"),
        ("`helmet`", "7.1", "security headers"),
        ("`cors`", "2.8", "restricts which origins may call the API"),
        ("`compression`", "1.7", "gzip responses"),
        ("`express-rate-limit`", "7.4", "the polling defence — see 14.6"),
        ("`pino` / `pino-http`", "9.3 / 10.2", "structured JSON logging"),
        ("`dotenv`", "16.4", "loads configuration from .env"),
        ("`zod`", "3.23", "request body validation"),
    ], widths=[52, 28, 85])
    para(doc,
         "Nine production dependencies and no test framework — the tests use Node 20's "
         "built-in test runner, which is why there is no jest or mocha here.")

    h2(doc, "2.4 Python dependencies")
    table(doc, ["Package", "Version", "What it does here"], [
        ("`flask`", "3.0.3", "the ML microservice HTTP layer"),
        ("`waitress`", "3.0.0", "production WSGI server, works well on Windows"),
        ("`scikit-learn`", "1.4.2", "the Random Forest itself"),
        ("`sentence-transformers`", "2.7.0", "all-MiniLM-L6-v2 embeddings, 384-d"),
        ("`torch`", "2.2.2", "CPU-only backend for the embedder"),
        ("`numpy` / `pandas`", "1.26 / 2.2", "arrays and dataframes"),
        ("`joblib`", "1.4.2", "saving and loading the trained model"),
        ("`psycopg2-binary`", "2.9.9", "direct PostgreSQL access for training"),
        ("`pgvector`", "0.2.5", "Python adapter for the vector type"),
        ("`ragas`", "0.1.9", "the four retrieval-quality metrics"),
        ("`langchain-groq`", "0.1.10", "connects RAGAS's judge to Groq"),
        ("`datasets`", "2.19.0", "required by RAGAS"),
    ], widths=[52, 28, 85])
    rich(doc, [
        ("Why the versions are pinned exactly. ", "b"),
        ("scikit-learn stores models in a version-sensitive format, and RAGAS 0.1.9 "
         "predates a LangChain reorganisation and imports symbols that later moved. "
         "Floating either would break silently, at scoring time, after a full evaluation "
         "run had already been generated — the worst possible moment.", ""),
    ])

    h2(doc, "2.5 Database extensions")
    table(doc, ["Extension", "What it provides", "Used by"], [
        ("`vector` (pgvector)", "the 384-dimension embedding column and cosine distance",
         "all document retrieval"),
        ("`pg_trgm`", "trigram similarity for fuzzy text matching",
         "faculty name matching"),
        ("`pgcrypto`", "UUID generation", "every primary key"),
    ], widths=[42, 78, 45])

    h2(doc, "2.6 The models")
    table(doc, ["Model", "Where it runs", "Job"], [
        ("`all-MiniLM-L6-v2`", "locally, in the Python service",
         "turns text into 384 numbers"),
        ("`RandomForestClassifier`", "locally, 300 trees",
         "predicts one of three availability states"),
        ("`openai/gpt-oss-120b`", "Groq, external", "writes the final sentence"),
        ("`openai/gpt-oss-20b`", "Groq, external", "judges answers during evaluation"),
    ], widths=[48, 52, 65])

    # ---------------------------------------------------------- ch 3
    h1(doc, "Chapter 3 — The repository, folder by folder")
    para(doc,
         "There are five top-level folders that matter. This chapter walks through every "
         "one of them and names every file, so you can open the project and know what "
         "you are looking at.", italic=True, color=GREY)

    h2(doc, "3.1 Top level")
    table(doc, ["Path", "What it is"], [
        ("`frontend/`", "the React website"),
        ("`backend/`", "the Express API server"),
        ("`database/`", "SQL schema, functions, security, migrations and seed data"),
        ("`machine-learning/`", "the Python service, training, ingestion, evaluation"),
        ("`notebooks/`", "generated Chapter 4 figures and tables"),
        ("`guide/`", "the scripts that built this document"),
        ("`start.bat` / `stop.bat`", "convenience scripts to launch all three processes"),
        ("`README.md`", "project overview"),
        ("`.gitignore`", "what git must never store — including .env"),
    ], widths=[52, 113])

    h2(doc, "3.2 backend/ — the API server")
    h3(doc, "backend/src/ — entry and configuration")
    filecard(doc, "src/index.js",
             "creates the Express app, applies middleware, starts listening on 4000, "
             "and runs startup diagnostics against the ML service")
    filecard(doc, "src/utilities/configuration.js",
             "reads every environment variable into one frozen config object; also "
             "defines PROMPT_TEMPLATE_VERSION and ROUTER_VERSION")
    filecard(doc, "src/utilities/service-clients.js",
             "the ONLY place the database, ML service and Groq are contacted")
    filecard(doc, "src/utilities/logger.js", "pino logger instance")
    filecard(doc, "src/utilities/ai-prompt-templates.js",
             "the single prompt template and the one conditional availability block")

    h3(doc, "backend/src/routes/ — the API surface")
    filecard(doc, "routes/index.js",
             "the public API: /chat, /map/pois, /faculty/search, /guard/*, /validate/*, "
             "/demo/*, /health, /eval/status, /me")
    filecard(doc, "routes/admin-routes.js",
             "/admin/pois (create, edit, unpublish, republish, delete, reindex, audit) "
             "and /admin/me/faculty for self-service visibility")

    h3(doc, "backend/src/services/ — the actual logic")
    filecard(doc, "services/knowledge-search-service.js",
             "runPipeline() — the orchestrator every question passes through")
    filecard(doc, "services/intent-query-router.js",
             "routeQuery() — decides the category and finds faculty candidates")
    filecard(doc, "services/faculty-presence-service.js",
             "getAvailability() — guard log, schedule, calendar, features, prediction")
    filecard(doc, "services/campus-places-service.js",
             "keeps the map and the retrieval corpus in step when a location changes")
    filecard(doc, "services/evaluation-runner.js",
             "the research harness — not part of any user request")

    h3(doc, "backend/src/middleware/")
    filecard(doc, "middleware/authentication.js",
             "requireAuth, requireRole and optionalAuth — verifies Supabase JWTs")
    filecard(doc, "middleware/privacy-masking-middleware.js",
             "maskPrediction, maskOverride, filterEgress, toChatDto — the privacy core")

    h3(doc, "backend/src/mock-services/ — DEMO_MODE")
    filecard(doc, "mock-services/mock-database.js",
             "an in-memory stand-in for every database call")
    filecard(doc, "mock-services/mock-campus-data.js", "placeholder locations and faculty")
    filecard(doc, "mock-services/mock-text-embeddings.js",
             "a lexical embedder, so retrieval runs with no Python service")
    filecard(doc, "mock-services/index.js", "chooses the mocks when DEMO_MODE is true")

    h3(doc, "backend/tests/ and scripts/")
    filecard(doc, "tests/campus-location-protocol.test.js",
             "the [LOCATION: id] tag and when the map may move")
    filecard(doc, "tests/chat-conversation.test.js",
             "conversation history sanitisation and prompt assembly")
    filecard(doc, "tests/privacy-masking-security.test.js",
             "five masking properties and the egress filter")
    filecard(doc, "tests/user-authorization.test.js",
             "role gates, the audit trail, and prompt-injection attempts")
    filecard(doc, "tests/temporal-consistency.test.js",
             "timezone boundaries and semester phase")
    filecard(doc, "tests/live-stack.test.js",
             "runs against the REAL Supabase, ML service and Groq")
    filecard(doc, "scripts/preflight-check.js",
             "one command that verifies the whole stack before a demo")

    h2(doc, "3.3 database/ — the schema and everything in it")
    filecard(doc, "tables-and-structure.sql",
             "all 25 tables, types, constraints and extensions. Opens by DROPPING the "
             "schema, so it runs on first install only")
    filecard(doc, "database-functions.sql",
             "the SQL functions: schedule_lookup_status, resolve_presence, "
             "match_document_chunks, resolve_faculty_candidates")
    filecard(doc, "security-and-permissions.sql",
             "row-level security, forced on every table, deny by default")
    filecard(doc, "apply-database-migrations.py",
             "applies the files above in the correct order; refuses to start if a "
             "migration on disk is missing from its list")
    h3(doc, "database/migrations/ — structural changes, applied once each")
    for name, what in [
        ("002_user_roles_and_locations.sql",
         "roles, the POI audit trail, faculty visibility columns"),
        ("003_campus_places_and_departments.sql", "the real campus locations"),
        ("004_map_pin_icons.sql", "per-location map icon"),
        ("005_schedule_campus.sql", "adds faculty_schedule.campus — the other-campus fix"),
        ("006_provision_accounts.sql", "admin and researcher role rows"),
        ("007_attendance_features.sql", "the attendance_features() function"),
        ("008_campus_aware_presence.sql",
         "teaches schedule_lookup_status about campuses"),
        ("009_dedupe_place_cards.sql", "removes duplicated place-card documents"),
    ]:
        filecard(doc, "migrations/" + name, what)
    h3(doc, "database/sample-data/ — reference and research data")
    for name, what in [
        ("001_sample_campus_records.sql", "SUPERSEDED placeholder data — do not run"),
        ("002_ccsict_schedule.sql", "the REAL timetable: 425 blocks, 37 lecturers"),
        ("003_synthetic_attendance.sql", "the generated attendance cohort SIM-01..37"),
        ("004_eval_queries.sql", "the pre-registered evaluation test set"),
        ("005_record_consent.sql", "TEMPLATE — run only against signed consent forms"),
        ("006_official_calendar.sql", "official ISU examination dates"),
        ("007_philippine_holidays.sql", "national holidays for 2026"),
    ]:
        filecard(doc, "sample-data/" + name, what)

    h2(doc, "3.4 machine-learning/ — the Python side")
    h3(doc, "The service")
    filecard(doc, "ai_api_service.py",
             "Flask on port 5001. Endpoints: /healthz, /model/info, /embed, "
             "/embed/batch, /tokenize/count, /predict")
    h3(doc, "Training")
    filecard(doc, "dataset_loader.py",
             "builds training samples from schedule, attendance and calendar; computes "
             "the historical rates strictly causally")
    filecard(doc, "feature_engineering.py",
             "defines the eleven features and turns a sample into a vector")
    filecard(doc, "train_availability_model.py",
             "trains the forest, refuses circular configurations, writes the artifact "
             "and the rf_model_version row")
    filecard(doc, "schedule_rule_baseline.py",
             "the rule the model is compared against")
    filecard(doc, "generate_synthetic_attendance.py",
             "produces the synthetic attendance cohort; reads the calendar rather than "
             "declaring it")
    h3(doc, "Documents and ingestion")
    filecard(doc, "convert_official_documents.py",
             "HTML or PDF to Markdown, preserving tables")
    filecard(doc, "document_knowledge_importer.py",
             "chunking, embedding and storage, with a checksum guard")
    filecard(doc, "institutional-documents/",
             "the converted School Calendar and Student Manual")
    h3(doc, "Support and evaluation")
    filecard(doc, "database_connector.py",
             "database access plus assert_research_ready(), the synthetic-data gate")
    filecard(doc, "schedule_importer.py", "loads the CCSICT timetable workbook")
    filecard(doc, "evaluate_rag_quality.py", "RAGAS scoring, one metric at a time")
    filecard(doc, "test_schedule_rule_contract.py",
             "proves the SQL and Python copies of the schedule rule agree")
    filecard(doc, "HISTORICAL-RESULTS.md",
             "frozen metrics from before the calendar was corrected")
    filecard(doc, "requirements.txt", "the pinned Python dependencies")

    h2(doc, "3.5 frontend/ — the website")
    h3(doc, "Entry and configuration")
    filecard(doc, "index.html · src/main.jsx · src/App.jsx",
             "page shell, React mount point, and the router")
    filecard(doc, "vite.config.js · tailwind.config.js · postcss.config.js",
             "build and styling configuration")
    filecard(doc, "src/styles/tokens.css · src/index.css", "design tokens and base CSS")
    h3(doc, "src/frontend-utilities/")
    filecard(doc, "backendApiClient.js",
             "THE single fetch wrapper. Every API call in the app goes through it")
    filecard(doc, "supabaseClient.js", "sign-in only — never used to read data")
    filecard(doc, "campusRoutingService.js · campusRoadGraph.json",
             "walking-route calculation across the campus path network")
    filecard(doc, "themeContext.jsx · appConstants.js", "dark mode and shared constants")
    h3(doc, "src/components/main-assistant/ — the primary surface")
    filecard(doc, "MainAssistantWorkspace.jsx", "holds the chat and map state together")
    filecard(doc, "ChatbotMessagePanel.jsx", "the conversation itself")
    filecard(doc, "InteractiveCampusMap.jsx", "Leaflet map, markers, routing overlay")
    filecard(doc, "PlaceDetailCard.jsx · PlacesDirectoryDrawer.jsx",
             "location detail and the browsable directory")
    filecard(doc, "FloatingChatDock.jsx", "the compact chat launcher")
    filecard(doc, "mapPinIconBuilder.js · mapMarkerGlyphs.js", "marker icon generation")
    h3(doc, "The other surfaces")
    filecard(doc, "components/admin-portal/",
             "AdminDashboard, CampusMapEditor, CampusLocationManager, "
             "AdminFacultyValidationPanel, AdminSecurityPresencePanel")
    filecard(doc, "components/faculty-validation-portal/",
             "FacultyAnswerValidationList, FacultyPrivacyToggleCard")
    filecard(doc, "components/security-guard-portal/",
             "SecurityGuardAttendanceDashboard")
    filecard(doc, "components/landing-page/",
             "fifteen components making up the marketing and explanation page")
    filecard(doc, "components/ui-primitives/ · shared-components/ · layout-patterns/",
             "buttons, inputs, banners, navigation, footers, page frames")
    filecard(doc, "custom-react-hooks/",
             "useCampusLocations, useReducedMotionPreference")


# =====================================================================
#  PART II — THE DATA FOUNDATION
# =====================================================================
def part2(doc):
    h1(doc, "PART II — The data foundation")
    para(doc, "Nothing else makes sense until you know what is stored and how it is "
              "shaped.", italic=True, color=GREY)

    h1(doc, "Chapter 4 — The database schema")
    para(doc,
         "One PostgreSQL database, one schema called `geobot`, %s tables. Row-level "
         "security is enabled AND forced on every one of them, with %s policies. "
         "Grouped by what they are for:" % (F["tables"], F["policies"]))

    h2(doc, "4.1 People")
    table(doc, ["Table", "Holds", "Notable columns"], [
        ("`faculty`", "one row per lecturer",
         "`is_consented`, `consent_date`, `availability_visible`, `data_origin`"),
        ("`faculty_alias`", "other names a lecturer is known by", "`alias`"),
        ("`faculty_pseudonym_map`", "the ONLY link between a person and attendance",
         "`faculty_id`, `pseudonym_id`"),
        ("`department`", "colleges and departments", "`name`"),
        ("`app_user_role`", "which login has which role",
         "`auth_user_id`, `role`, `faculty_id`"),
        ("`guard_user`", "security accounts", "—"),
    ], widths=[44, 60, 61])

    h2(doc, "4.2 Time and place")
    table(doc, ["Table", "Holds", "Notable columns"], [
        ("`faculty_schedule`", "the timetable",
         "`day_of_week`, `start_time`, `end_time`, `block_kind`, `campus`"),
        ("`institutional_event`", "holidays, exam periods, semester markers",
         "`disrupts_schedule`"),
        ("`poi`", "campus locations",
         "`slug`, `lat`, `lng`, `poi_type`, `is_published`, `icon`"),
        ("`poi_audit`", "every change made to a location", "`action`, `before_state`"),
    ], widths=[44, 60, 61])

    h2(doc, "4.3 Evidence about presence")
    table(doc, ["Table", "Holds", "Notable columns"], [
        ("`attendance_record`", "biometric-style punches",
         "`pseudonym_id` — NO faculty_id"),
        ("`guard_presence_event`", "arrivals and departures logged by a guard",
         "`event_type`, `occurred_at`"),
        ("`availability_status`", "the three status codes and their labels",
         "`code`, `display_label`"),
        ("`faculty_validation`", "a lecturer confirming whether an estimate was right",
         "`system_status`, `actual_status`, `correctness`"),
        ("`faculty_visibility_event`", "a record of pausing or resuming visibility", "—"),
    ], widths=[44, 60, 61])

    h2(doc, "4.4 The retrieval corpus")
    table(doc, ["Table", "Holds", "Notable columns"], [
        ("`document`", "one row per source document",
         "`doc_type`, `source_origin`, `source_checksum`"),
        ("`document_chunk`", "the searchable pieces",
         "`content`, `embedding vector(384)`, `token_count`, `poi_id`"),
        ("`poi_document`", "generated place-card text per location", "`generated_text`"),
    ], widths=[44, 60, 61])

    h2(doc, "4.5 Research and operations")
    table(doc, ["Table", "Holds"], [
        ("`eval_query`", "the pre-registered test set with ground-truth answers"),
        ("`eval_run`", "one evaluation run: model, judge, prompt version, settings"),
        ("`eval_result`", "one answer: retrieved contexts, the answer, all timings"),
        ("`ragas_score`", "the four quality metrics per result"),
        ("`rf_model_version`", "every trained model: metrics, features, data_origin"),
        ("`chat_log`", "live traffic — kept SEPARATE from the evaluation tables"),
        ("`demo_query`", "canned comparison queries for the public demo"),
    ], widths=[46, 119])
    rich(doc, [
        ("Why chat_log is separate. ", "b"),
        ("If development traffic and evaluation results shared a table, someone would "
         "eventually report one as the other. Separate tables make that mistake "
         "impossible rather than merely discouraged.", ""),
    ])

    h2(doc, "4.6 The one column on almost every table")
    para(doc,
         "`data_origin` is either `real` or `synthetic`. It appears on faculty, "
         "schedules, attendance, documents, chunks, locations, events, evaluation rows "
         "and model versions. It is the mechanism that stops invented data being "
         "reported as a finding, and it is checked by a database function called "
         "`corpus_is_research_ready()` before any metric may be produced.")

    h1(doc, "Chapter 5 — The database functions")
    para(doc,
         "Logic that lives in SQL rather than JavaScript, because it must be identical "
         "for the live system and for training.")
    table(doc, ["Function", "Answers"], [
        ("`schedule_lookup_status(faculty, at, semester, tz, campus)`",
         "what does the timetable say about this person at this moment on this campus?"),
        ("`resolve_presence(faculty, at, tz)`",
         "has a guard seen them arrive or leave today?"),
        ("`match_document_chunks(embedding, k, floor, include_synthetic)`",
         "which document chunks are closest to this question?"),
        ("`resolve_faculty_candidates(needle, limit)`",
         "which consented lecturers might this name refer to?"),
        ("`faculty_is_answerable(faculty)`",
         "may we say anything at all about this person?"),
        ("`attendance_features(pseudonym, at, tz)`",
         "what are this person's historical presence rates?"),
        ("`corpus_is_research_ready()`",
         "which tables still contain synthetic rows?"),
        ("`has_role(role)` / `validator_faculty_id()`",
         "security helpers used by the row-level policies"),
    ], widths=[74, 91])

    h1(doc, "Chapter 6 — What is actually in the database")
    para(doc, "Counts read from the live system while this document was built.",
         italic=True, color=GREY)
    table(doc, ["Table", "Rows", "What that means"], [
        ("`poi`", str(F["poi"]), "%s published campus locations" % F["poi_published"]),
        ("`document`", str(F["document"]),
         "the Calendar, the Student Manual, and one place-card per location"),
        ("`document_chunk`", str(F["document_chunk"]), "the searchable corpus"),
        ("`faculty`", str(F["faculty_total"]),
         "%s real lecturers plus a synthetic cohort" % F["faculty_real"]),
        ("`faculty_schedule`", str(F["schedule_total"]),
         "%s class blocks and %s consultation blocks"
         % (F["schedule_class"], F["schedule_consult"])),
        ("`attendance_record`", str(F["attendance"]), "ALL synthetic — see Chapter 24"),
        ("`institutional_event`", str(F["institutional_event"]),
         "holidays, exam periods, semester markers"),
        ("`guard_presence_event`", str(F["guard_events"]), "test data — see Chapter 24"),
        ("`faculty_validation`", str(F["faculty_validation"]), "real validator entries"),
        ("`eval_query`", str(F["eval_query"]), "the registered test set"),
        ("`eval_result`", str(F["eval_result"]), "two runs, both arms"),
        ("`ragas_score`", str(F["ragas_score"]), "not yet scored — see Chapter 24"),
    ], widths=[46, 22, 97])
